from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from fastapi.responses import StreamingResponse
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from clauseai_backend.api.deps import get_current_user, get_reference_db, get_user_db
from clauseai_backend.core.config import settings
from clauseai_backend.models.auth import User
from clauseai_backend.models.projects import (
    AnalysisArtifact,
    BillDraft,
    BillDraftVersion,
    Project,
    ProjectMetadata,
    SourceDocument,
    Suggestion,
)
from clauseai_backend.schemas.projects import (
    ArtifactResponse,
    AgentPassResponse,
    DraftResponse,
    DraftSaveRequest,
    MetadataUpdateRequest,
    DraftVersionResponse,
    MetadataResponse,
    SuggestionActionRequest,
    SuggestionResponse,
)
from clauseai_backend.services.analysis import (
    AnalysisResult,
    build_legal_analysis,
    build_similar_bills_analysis,
    build_stakeholder_analysis,
)
from clauseai_backend.services.drafting_agent import build_editor_agent_pass
from clauseai_backend.services.file_extraction import detect_file_type, extract_text_from_bytes
from clauseai_backend.services.metadata import generate_metadata_from_text

router = APIRouter(prefix="/api/projects", tags=["workflow-content"])


def _get_project_or_404(db: Session, project_id: str, user_id: str) -> Project:
    project = db.scalar(select(Project).where(Project.project_id == project_id, Project.user_id == user_id))
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    return project


def _get_draft_or_404(db: Session, project_id: str) -> BillDraft:
    draft = db.scalar(select(BillDraft).where(BillDraft.project_id == project_id))
    if not draft:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Draft not found")
    return draft


def _upsert_metadata(db: Session, project: Project, payload: dict[str, object]) -> ProjectMetadata:
    record = db.get(ProjectMetadata, project.project_id)
    if not record:
        record = ProjectMetadata(
            project_id=project.project_id,
            title=str(payload["title"]),
            description=str(payload["description"]),
            summary=str(payload["summary"]),
            keywords=list(payload["keywords"]),
            generated_json=payload,
        )
        db.add(record)
    else:
        record.title = str(payload["title"])
        record.description = str(payload["description"])
        record.summary = str(payload["summary"])
        record.keywords = list(payload["keywords"])
        record.generated_json = payload
        db.add(record)
    project.title = record.title
    db.add(project)
    db.flush()
    return record


def _replace_stage_outputs(db: Session, project_id: str, stage_name: str, result: AnalysisResult) -> tuple[AnalysisArtifact, list[Suggestion]]:
    db.execute(delete(Suggestion).where(Suggestion.project_id == project_id, Suggestion.stage_name == stage_name))
    db.execute(delete(AnalysisArtifact).where(AnalysisArtifact.project_id == project_id, AnalysisArtifact.stage_name == stage_name))
    artifact = AnalysisArtifact(
        artifact_id=str(uuid4()),
        project_id=project_id,
        stage_name=stage_name,
        artifact_kind="report",
        status="ready",
        markdown_content=result.markdown,
        payload_json=result.payload,
    )
    db.add(artifact)
    suggestions: list[Suggestion] = []
    for item in result.suggestions:
        suggestion = Suggestion(
            suggestion_id=str(uuid4()),
            project_id=project_id,
            stage_name=stage_name,
            title=str(item["title"]),
            rationale=str(item["rationale"]),
            before_text=str(item.get("before_text") or ""),
            after_text=str(item.get("after_text") or ""),
            source_refs=list(item.get("source_refs") or []),
        )
        db.add(suggestion)
        suggestions.append(suggestion)
    db.flush()
    return artifact, suggestions


def _serialize_suggestion(item: Suggestion) -> SuggestionResponse:
    return SuggestionResponse(
        suggestion_id=item.suggestion_id,
        stage_name=item.stage_name,
        title=item.title,
        rationale=item.rationale,
        before_text=item.before_text,
        after_text=item.after_text,
        source_refs=list(item.source_refs),
        status=item.status,
    )


def _store_draft_version(
    db: Session,
    *,
    project: Project,
    draft: BillDraft,
    content_text: str,
    source_kind: str,
    reason: str,
    created_by: str,
    metadata: dict[str, object] | None = None,
) -> BillDraftVersion:
    draft.current_text = content_text
    db.add(draft)
    version = BillDraftVersion(
        version_id=str(uuid4()),
        draft_id=draft.draft_id,
        project_id=project.project_id,
        version_number=_next_version_number(db, draft.project_id),
        source_kind=source_kind,
        content_text=content_text,
        change_summary={"reason": reason, **(metadata or {})},
        created_by=created_by,
    )
    db.add(version)
    db.flush()
    return version


@router.post("/{project_id}/source-document", status_code=status.HTTP_201_CREATED)
async def upload_source_document(
    project_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> dict[str, object]:
    project = _get_project_or_404(db, project_id, current_user.user_id)
    file_bytes = await file.read()
    file_type = detect_file_type(file.filename or "upload.txt")
    extracted_text = extract_text_from_bytes(file_type, file_bytes)

    storage_dir = settings.storage_root / project_id
    storage_dir.mkdir(parents=True, exist_ok=True)
    storage_path = storage_dir / f"{uuid4()}-{Path(file.filename or 'upload').name}"
    storage_path.write_bytes(file_bytes)

    document = SourceDocument(
        document_id=str(uuid4()),
        project_id=project.project_id,
        filename=file.filename or "upload",
        file_type=file_type,
        storage_path=str(storage_path),
        extracted_text=extracted_text,
    )
    draft = _get_draft_or_404(db, project.project_id)
    _store_draft_version(
        db,
        project=project,
        draft=draft,
        content_text=extracted_text,
        source_kind="upload_extract",
        reason="Source document uploaded and extracted",
        created_by=current_user.user_id,
    )

    project.current_stage = "extraction"
    db.add(project)
    db.add(document)
    db.commit()
    return {
        "document_id": document.document_id,
        "filename": document.filename,
        "file_type": document.file_type,
        "extracted_text": extracted_text,
    }


def _next_version_number(db: Session, project_id: str) -> int:
    versions = db.execute(
        select(BillDraftVersion.version_number).where(BillDraftVersion.project_id == project_id)
    ).scalars()
    current_max = max(list(versions) or [0])
    return current_max + 1


@router.get("/{project_id}/source-document/latest")
def get_latest_source_document(
    project_id: str,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> dict[str, object]:
    _get_project_or_404(db, project_id, current_user.user_id)
    document = db.scalar(
        select(SourceDocument).where(SourceDocument.project_id == project_id).order_by(SourceDocument.created_at.desc())
    )
    if not document:
        return {
            "document_id": None,
            "filename": None,
            "file_type": None,
            "extracted_text": "",
        }
    return {
        "document_id": document.document_id,
        "filename": document.filename,
        "file_type": document.file_type,
        "extracted_text": document.extracted_text,
    }


@router.post("/{project_id}/metadata/generate", response_model=MetadataResponse)
def generate_metadata(
    project_id: str,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> MetadataResponse:
    project = _get_project_or_404(db, project_id, current_user.user_id)
    draft = _get_draft_or_404(db, project.project_id)
    if not draft or not draft.current_text.strip():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Upload and extract a source document first")

    payload = generate_metadata_from_text(draft.current_text, project.title)
    metadata = _upsert_metadata(db, project, payload)
    project.current_stage = "metadata"
    db.add(project)
    db.commit()
    return MetadataResponse(
        title=metadata.title,
        description=metadata.description,
        summary=metadata.summary,
        keywords=list(metadata.keywords),
        extras={key: value for key, value in metadata.generated_json.items() if key not in {"title", "description", "summary", "keywords"}},
    )


@router.get("/{project_id}/metadata", response_model=MetadataResponse)
def get_metadata(
    project_id: str,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> MetadataResponse:
    project = _get_project_or_404(db, project_id, current_user.user_id)
    metadata = db.get(ProjectMetadata, project_id)
    if not metadata:
        return MetadataResponse(
            title=project.title,
            description="",
            summary="",
            keywords=[],
            extras={},
        )
    return MetadataResponse(
        title=metadata.title,
        description=metadata.description,
        summary=metadata.summary,
        keywords=list(metadata.keywords),
        extras={key: value for key, value in metadata.generated_json.items() if key not in {"title", "description", "summary", "keywords"}},
    )


@router.put("/{project_id}/metadata", response_model=MetadataResponse)
def update_metadata(
    project_id: str,
    payload: MetadataUpdateRequest,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> MetadataResponse:
    project = _get_project_or_404(db, project_id, current_user.user_id)
    metadata = db.get(ProjectMetadata, project_id)
    if not metadata:
        metadata = ProjectMetadata(
            project_id=project_id,
            title=payload.title,
            description=payload.description,
            summary=payload.summary,
            keywords=list(payload.keywords),
            generated_json={},
        )
    generated_json = {
        "title": payload.title,
        "description": payload.description,
        "summary": payload.summary,
        "keywords": list(payload.keywords),
        **payload.extras,
    }
    metadata.title = payload.title
    metadata.description = payload.description
    metadata.summary = payload.summary
    metadata.keywords = list(payload.keywords)
    metadata.generated_json = generated_json
    project.title = payload.title
    project.current_stage = "metadata"
    db.add(metadata)
    db.add(project)
    db.commit()
    return MetadataResponse(
        title=metadata.title,
        description=metadata.description,
        summary=metadata.summary,
        keywords=list(metadata.keywords),
        extras=payload.extras,
    )


@router.post("/{project_id}/analysis/{stage_name}", response_model=ArtifactResponse)
def generate_stage_analysis(
    project_id: str,
    stage_name: str,
    user_db: Session = Depends(get_user_db),
    reference_db: Session = Depends(get_reference_db),
    current_user: User = Depends(get_current_user),
) -> ArtifactResponse:
    project = _get_project_or_404(user_db, project_id, current_user.user_id)
    metadata = user_db.get(ProjectMetadata, project.project_id)
    if not metadata:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Metadata must be generated first")

    if stage_name == "similar-bills":
        result = build_similar_bills_analysis(
            reference_db,
            title=metadata.title,
            summary=metadata.summary,
            keywords=list(metadata.keywords),
            draft_text=_get_draft_or_404(user_db, project.project_id).current_text,
        )
    elif stage_name == "legal":
        result = build_legal_analysis(
            reference_db,
            title=metadata.title,
            summary=metadata.summary,
            keywords=list(metadata.keywords),
            draft_text=_get_draft_or_404(user_db, project.project_id).current_text,
        )
    elif stage_name == "stakeholders":
        result = build_stakeholder_analysis(
            title=metadata.title,
            summary=metadata.summary,
            keywords=list(metadata.keywords),
            draft_text=_get_draft_or_404(user_db, project.project_id).current_text,
        )
    else:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Unknown analysis stage")

    artifact, _ = _replace_stage_outputs(user_db, project.project_id, stage_name, result)
    project.current_stage = stage_name
    user_db.add(project)
    user_db.commit()
    return ArtifactResponse(
        artifact_id=artifact.artifact_id,
        stage_name=artifact.stage_name,
        artifact_kind=artifact.artifact_kind,
        status=artifact.status,
        markdown_content=artifact.markdown_content,
        payload_json=artifact.payload_json,
    )


@router.get("/{project_id}/analysis/{stage_name}", response_model=ArtifactResponse)
def get_stage_analysis(
    project_id: str,
    stage_name: str,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> ArtifactResponse:
    _get_project_or_404(db, project_id, current_user.user_id)
    artifact = db.scalar(
        select(AnalysisArtifact)
        .where(AnalysisArtifact.project_id == project_id, AnalysisArtifact.stage_name == stage_name)
        .order_by(AnalysisArtifact.updated_at.desc())
    )
    if not artifact:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Analysis not generated yet")
    return ArtifactResponse(
        artifact_id=artifact.artifact_id,
        stage_name=artifact.stage_name,
        artifact_kind=artifact.artifact_kind,
        status=artifact.status,
        markdown_content=artifact.markdown_content,
        payload_json=artifact.payload_json,
    )


@router.get("/{project_id}/suggestions/{stage_name}", response_model=list[SuggestionResponse])
def list_stage_suggestions(
    project_id: str,
    stage_name: str,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> list[SuggestionResponse]:
    _get_project_or_404(db, project_id, current_user.user_id)
    suggestions = (
        db.execute(
            select(Suggestion)
            .where(Suggestion.project_id == project_id, Suggestion.stage_name == stage_name)
            .order_by(Suggestion.created_at.asc())
        )
        .scalars()
        .all()
    )
    return [_serialize_suggestion(item) for item in suggestions]


@router.get("/{project_id}/suggestions", response_model=list[SuggestionResponse])
def list_all_suggestions(
    project_id: str,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> list[SuggestionResponse]:
    _get_project_or_404(db, project_id, current_user.user_id)
    suggestions = (
        db.execute(
            select(Suggestion).where(Suggestion.project_id == project_id).order_by(Suggestion.created_at.asc())
        )
        .scalars()
        .all()
    )
    return [_serialize_suggestion(item) for item in suggestions]


@router.get("/{project_id}/draft", response_model=DraftResponse)
def get_current_draft(
    project_id: str,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> DraftResponse:
    _get_project_or_404(db, project_id, current_user.user_id)
    draft = _get_draft_or_404(db, project_id)
    return DraftResponse(draft_id=draft.draft_id, current_text=draft.current_text, title=draft.title)


@router.get("/{project_id}/draft/versions", response_model=list[DraftVersionResponse])
def list_draft_versions(
    project_id: str,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> list[DraftVersionResponse]:
    _get_project_or_404(db, project_id, current_user.user_id)
    versions = (
        db.execute(
            select(BillDraftVersion)
            .where(BillDraftVersion.project_id == project_id)
            .order_by(BillDraftVersion.version_number.desc())
        )
        .scalars()
        .all()
    )
    return [
        DraftVersionResponse(
            version_id=item.version_id,
            version_number=item.version_number,
            source_kind=item.source_kind,
            content_text=item.content_text,
            change_summary=item.change_summary,
            created_at=item.created_at,
        )
        for item in versions
    ]


@router.post("/{project_id}/draft/versions", response_model=DraftVersionResponse, status_code=status.HTTP_201_CREATED)
def save_draft_version(
    project_id: str,
    payload: DraftSaveRequest,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> DraftVersionResponse:
    project = _get_project_or_404(db, project_id, current_user.user_id)
    draft = _get_draft_or_404(db, project_id)
    version = _store_draft_version(
        db,
        project=project,
        draft=draft,
        content_text=payload.content_text,
        source_kind="manual_edit",
        reason=payload.change_reason,
        created_by=current_user.user_id,
    )
    project.current_stage = "editor"
    project.status = "in_review"
    db.add(project)
    db.commit()
    return DraftVersionResponse(
        version_id=version.version_id,
        version_number=version.version_number,
        source_kind=version.source_kind,
        content_text=version.content_text,
        change_summary=version.change_summary,
        created_at=version.created_at,
    )


@router.post("/{project_id}/draft/versions/{version_id}/restore", response_model=DraftVersionResponse)
def restore_draft_version(
    project_id: str,
    version_id: str,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> DraftVersionResponse:
    project = _get_project_or_404(db, project_id, current_user.user_id)
    draft = _get_draft_or_404(db, project_id)
    selected = db.scalar(
        select(BillDraftVersion).where(
            BillDraftVersion.project_id == project_id, BillDraftVersion.version_id == version_id
        )
    )
    if not selected:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Version not found")
    version = _store_draft_version(
        db,
        project=project,
        draft=draft,
        content_text=selected.content_text,
        source_kind="restore",
        reason=f"Restored version {selected.version_number}",
        created_by=current_user.user_id,
        metadata={"restored_version_id": selected.version_id},
    )
    project.current_stage = "editor"
    db.add(project)
    db.commit()
    return DraftVersionResponse(
        version_id=version.version_id,
        version_number=version.version_number,
        source_kind=version.source_kind,
        content_text=version.content_text,
        change_summary=version.change_summary,
        created_at=version.created_at,
    )


@router.post("/{project_id}/suggestion-items/{suggestion_id}/apply", response_model=SuggestionResponse)
def apply_suggestion(
    project_id: str,
    suggestion_id: str,
    payload: SuggestionActionRequest,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> SuggestionResponse:
    project = _get_project_or_404(db, project_id, current_user.user_id)
    draft = _get_draft_or_404(db, project_id)
    suggestion = db.scalar(
        select(Suggestion).where(Suggestion.project_id == project_id, Suggestion.suggestion_id == suggestion_id)
    )
    if not suggestion:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Suggestion not found")

    replacement_text = (payload.after_text or suggestion.after_text).strip()
    if not replacement_text:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Suggestion has no replacement text")

    if suggestion.before_text and suggestion.before_text in draft.current_text:
        next_text = draft.current_text.replace(suggestion.before_text, replacement_text, 1)
    elif replacement_text in draft.current_text:
        next_text = draft.current_text
    else:
        separator = "\n\n" if draft.current_text.strip() else ""
        next_text = f"{draft.current_text.rstrip()}{separator}{replacement_text}".strip()

    _store_draft_version(
        db,
        project=project,
        draft=draft,
        content_text=next_text,
        source_kind="suggestion_accept",
        reason=payload.change_reason,
        created_by=current_user.user_id,
        metadata={"suggestion_id": suggestion.suggestion_id, "stage_name": suggestion.stage_name},
    )
    suggestion.status = "accepted"
    db.add(suggestion)
    project.current_stage = "editor"
    project.status = "in_review"
    db.add(project)
    db.commit()
    return _serialize_suggestion(suggestion)


@router.post("/{project_id}/suggestion-items/{suggestion_id}/reject", response_model=SuggestionResponse)
def reject_suggestion(
    project_id: str,
    suggestion_id: str,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> SuggestionResponse:
    _get_project_or_404(db, project_id, current_user.user_id)
    suggestion = db.scalar(
        select(Suggestion).where(Suggestion.project_id == project_id, Suggestion.suggestion_id == suggestion_id)
    )
    if not suggestion:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Suggestion not found")
    suggestion.status = "rejected"
    db.add(suggestion)
    db.commit()
    return _serialize_suggestion(suggestion)


@router.post("/{project_id}/editor/agent-pass", response_model=AgentPassResponse)
def run_editor_agent_pass(
    project_id: str,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> AgentPassResponse:
    project = _get_project_or_404(db, project_id, current_user.user_id)
    draft = _get_draft_or_404(db, project_id)
    metadata = db.get(ProjectMetadata, project_id)
    suggestion_rows = (
        db.execute(
            select(Suggestion).where(Suggestion.project_id == project_id, Suggestion.status == "pending").order_by(Suggestion.created_at.asc())
        )
        .scalars()
        .all()
    )
    payload = build_editor_agent_pass(
        draft_text=draft.current_text,
        project_title=project.title,
        metadata=metadata.generated_json if metadata else {},
        suggestions=[
            {
                "stage_name": item.stage_name,
                "title": item.title,
                "rationale": item.rationale,
                "before_text": item.before_text,
                "after_text": item.after_text,
                "source_refs": list(item.source_refs),
            }
            for item in suggestion_rows
        ],
    )
    if not payload:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Codex runtime is not available")

    result = AnalysisResult(
        markdown=str(payload.get("report") or "No report generated."),
        payload=payload,
        suggestions=[
            {
                "title": str(item.get("title") or "Untitled suggestion"),
                "rationale": str(item.get("rationale") or ""),
                "before_text": str(item.get("before_text") or ""),
                "after_text": str(item.get("after_text") or ""),
                "source_refs": list(item.get("source_refs") or []),
            }
            for item in (payload.get("suggestions") or [])
            if isinstance(item, dict)
        ],
    )
    artifact, suggestions = _replace_stage_outputs(db, project.project_id, "editor", result)
    project.current_stage = "editor"
    project.status = "in_review"
    db.add(project)
    db.commit()
    return AgentPassResponse(
        artifact_id=artifact.artifact_id,
        markdown_content=artifact.markdown_content,
        payload_json=artifact.payload_json,
        suggestion_count=len(suggestions),
    )


@router.get("/{project_id}/export/txt")
def export_draft_txt(
    project_id: str,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> StreamingResponse:
    project = _get_project_or_404(db, project_id, current_user.user_id)
    draft = _get_draft_or_404(db, project_id)
    filename = f"{_slugify(project.title)}.txt"
    return StreamingResponse(
        iter([draft.current_text.encode("utf-8")]),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{project_id}/export/docx")
def export_draft_docx(
    project_id: str,
    db: Session = Depends(get_user_db),
    current_user: User = Depends(get_current_user),
) -> StreamingResponse:
    project = _get_project_or_404(db, project_id, current_user.user_id)
    draft = _get_draft_or_404(db, project_id)
    document = Document()
    document.add_heading(project.title, level=1)
    for paragraph in draft.current_text.split("\n\n"):
        cleaned = paragraph.strip()
        if cleaned:
            document.add_paragraph(cleaned)
    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    filename = f"{_slugify(project.title)}.docx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _slugify(value: str) -> str:
    cleaned = "".join(ch.lower() if ch.isalnum() else "-" for ch in value).strip("-")
    return cleaned or "draft"
